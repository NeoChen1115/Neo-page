#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def read_markdown(filepath):
    """Read markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown(content):
    """Parse markdown content into sections"""
    lines = content.split('\n')
    sections = []
    current_section = {'type': 'paragraph', 'content': []}
    
    for line in lines:
        if line.startswith('# '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'title', 'content': line[2:].strip()}
            sections.append(current_section)
            current_section = {'type': 'paragraph', 'content': []}
        elif line.startswith('## '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'heading2', 'content': line[3:].strip()}
            sections.append(current_section)
            current_section = {'type': 'paragraph', 'content': []}
        elif line.startswith('### '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'heading3', 'content': line[4:].strip()}
            sections.append(current_section)
            current_section = {'type': 'paragraph', 'content': []}
        elif line.startswith('| '):
            if current_section['content']:
                sections.append(current_section)
            # Parse table
            table_lines = [line]
            idx = lines.index(line)
            while idx + 1 < len(lines) and lines[idx + 1].startswith('| '):
                idx += 1
                table_lines.append(lines[idx])
            current_section = {'type': 'table', 'content': table_lines}
            sections.append(current_section)
            current_section = {'type': 'paragraph', 'content': []}
        elif line.startswith('- ') or line.startswith('* '):
            current_section['type'] = 'bullet'
            current_section['content'].append(line[2:].strip())
        elif line.startswith('[') and line.endswith(']'):
            # References
            if current_section['content']:
                sections.append(current_section)
            current_section = {'type': 'reference', 'content': line.strip()}
            sections.append(current_section)
            current_section = {'type': 'paragraph', 'content': []}
        elif line.strip() == '---':
            continue
        elif line.strip():
            current_section['content'].append(line)
    
    if current_section['content']:
        sections.append(current_section)
    
    return sections

def add_sections_to_doc(doc, sections):
    """Add parsed sections to Word document"""
    for section in sections:
        if section['type'] == 'title':
            p = doc.add_heading(section['content'], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif section['type'] == 'heading2':
            doc.add_heading(section['content'], level=1)
        elif section['type'] == 'heading3':
            doc.add_heading(section['content'], level=2)
        elif section['type'] == 'bullet':
            for item in section['content']:
                if item:
                    doc.add_paragraph(item, style='List Bullet')
        elif section['type'] == 'table':
            parse_and_add_table(doc, section['content'])
        elif section['type'] == 'reference':
            p = doc.add_paragraph(section['content'])
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
        elif section['type'] == 'paragraph':
            for line in section['content']:
                if line.strip():
                    # Handle bold and italic
                    p = doc.add_paragraph()
                    add_formatted_text(p, line)

def add_formatted_text(paragraph, text):
    """Add text with formatting (bold, italic) to paragraph"""
    # Simple pattern matching for ** (bold) and * (italic)
    pattern = r'\*\*([^\*]+)\*\*|\*([^\*]+)\*|__([^_]+)__|_([^_]+)_'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add normal text before match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Add formatted text
        matched_text = match.group(1) or match.group(2) or match.group(3) or match.group(4)
        run = paragraph.add_run(matched_text)
        
        if match.group(1) or match.group(3):  # Bold
            run.bold = True
        elif match.group(2) or match.group(4):  # Italic
            run.italic = True
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def parse_and_add_table(doc, table_lines):
    """Parse and add table to document"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    header_line = table_lines[0]
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    
    # Skip separator line
    # Parse rows
    rows = []
    for line in table_lines[2:]:
        if line.startswith('| '):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    
    # Create table
    if rows:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add header
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Add rows
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell in enumerate(row):
                row_cells[j].text = cell

def markdown_to_word(md_file, word_file, title=None, author=None):
    """Convert markdown file to Word document"""
    # Read markdown
    content = read_markdown(md_file)
    
    # Create Word document
    doc = Document()
    
    # Add title page info if provided
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
    
    if author:
        author_para = doc.add_paragraph(f"Author: {author}")
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Parse and add content
    sections = parse_markdown(content)
    add_sections_to_doc(doc, sections)
    
    # Save
    doc.save(word_file)
    print(f"✅ Generated: {word_file}")

# Generate Word files
print("Generating Word files...\n")

# Chinese simple version
markdown_to_word(
    '/workspaces/Neo-page/Survey_RAG_Simple_Version.md',
    '/workspaces/Neo-page/Survey_RAG_Simple_Version.docx',
    title='Retrieval-Augmented Generation for Big Data Survey',
    author='Student'
)

# English simple version
markdown_to_word(
    '/workspaces/Neo-page/Survey_RAG_Simple_Version_EN.md',
    '/workspaces/Neo-page/Survey_RAG_Simple_Version_EN.docx',
    title='Retrieval-Augmented Generation for Big Data: A Simple Survey',
    author='Student'
)

print("\n✨ All Word files generated successfully!")
